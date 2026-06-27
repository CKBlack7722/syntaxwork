from __future__ import annotations

import argparse
import csv
import json
import re
import shutil
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path

import openpyxl
from docx import Document

from refresh_logic_group_from_docx import (
    cell_text,
    expand_codes,
    load_all_vars,
    natural_key,
    norm_qid,
    option_expr,
    parse_condition,
    render_node,
    target_vars_for_qid,
)


LOGIC_SHEET = "邏輯組"
MUTEX_HEADER = "互斥"


@dataclass(frozen=True)
class MutexRule:
    qid: str
    option_code: str
    condition: str
    mutex: str
    bracket: str
    source_text: str


@dataclass(frozen=True)
class TableResolution:
    """A table-to-question match that is safe enough to use automatically."""

    qid: str
    option_codes: tuple[str, ...]
    candidates: tuple[str, ...]
    reason: str
    table_index: int


@dataclass(frozen=True)
class TextUnit:
    text: str
    table: TableResolution | None = None


def normalize(value: str) -> str:
    text = unicodedata.normalize("NFKC", value or "")
    text = text.replace("～", "~").replace("−", "-")
    text = re.sub(r"\s+", "", text)
    return text


def multi_groups(all_vars: set[str]) -> dict[str, set[str]]:
    groups: dict[str, set[str]] = {}
    for name in all_vars:
        match = re.fullmatch(r"(v.+)m(\d+)", name, flags=re.I)
        if match:
            groups.setdefault(match.group(1), set()).add(str(int(match.group(2))))
    return groups


def infer_table_qid(cells: list[str], groups: dict[str, set[str]], table_index: int) -> TableResolution:
    """Resolve a table only when its option set identifies one question uniquely.

    A Word table number is intentionally not part of this rule. Table numbers move
    whenever a questionnaire is edited, while a complete option-code set is useful
    evidence. A partial overlap is never enough: it is reported for review instead.
    """
    ignored_codes = {"96", "97", "98", "99"}
    codes = {
        str(int(code))
        for cell in cells
        for code in re.findall(r"\((\d+)\)", unicodedata.normalize("NFKC", cell))
    } - ignored_codes
    option_codes = tuple(sorted(codes, key=int))
    if not codes:
        return TableResolution("", option_codes, (), "table has no usable option codes", table_index)

    exact = sorted(
        name[1:]
        for name, group_codes in groups.items()
        if (group_codes - ignored_codes) == codes
    )
    if len(exact) == 1:
        return TableResolution(exact[0], option_codes, tuple(exact), "unique exact option-set match", table_index)
    if exact:
        return TableResolution(
            "",
            option_codes,
            tuple(exact),
            "ambiguous exact option-set match; user confirmation is required",
            table_index,
        )

    overlaps = sorted(
        name[1:]
        for name, group_codes in groups.items()
        if codes & (group_codes - ignored_codes)
    )
    return TableResolution(
        "",
        option_codes,
        tuple(overlaps),
        "no unique exact option-set match; user confirmation is required",
        table_index,
    )


def text_units(docx_path: Path, all_vars: set[str]) -> list[TextUnit]:
    doc = Document(docx_path)
    units: list[TextUnit] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            units.append(TextUnit(text))
    groups = multi_groups(all_vars)
    for table_index, table in enumerate(doc.tables, start=1):
        cells = [cell.text.strip().replace("\n", " ") for row in table.rows for cell in row.cells if cell.text.strip()]
        hint = infer_table_qid(cells, groups, table_index)
        for cell in cells:
            # Keep every cell separate: a row may contain option 05, 16 and 27.
            units.append(TextUnit(cell, hint))
    return units


def qid_from_question(text: str) -> str:
    clean = unicodedata.normalize("NFKC", text.strip())
    clean = re.sub(r"^[\s*◎●○□壹貳參一二三四五六七八九十、.()（）-]+", "", clean)
    match = re.match(r"(?:CK)?([A-Z][A-Z0-9_]{0,12})(?=[^A-Z0-9_]|$)", clean)
    return norm_qid(match.group(1)) if match else ""


def option_code_from_text(text: str) -> str:
    match = re.search(r"(?:^|[□\s])\((\d+)\)", unicodedata.normalize("NFKC", text))
    return str(int(match.group(1))) if match else ""


def expr_for_option(qid: str, code: str, all_vars: set[str]) -> str:
    node = option_expr(qid, code, all_vars)
    return render_node(node) if node else ""


def expr_for_existing_options(qid: str, codes: list[str], all_vars: set[str]) -> str:
    exprs: list[str] = []
    for code in codes:
        expr = expr_for_option(qid, code, all_vars)
        if expr and expr not in exprs:
            exprs.append(expr)
    return " | ".join(exprs)


def parse_option_range_target(text: str, current_qid: str, all_vars: set[str]) -> str:
    match = re.search(r"選項\(?(\d+)\)?(?:\s*[~\-至]\s*\(?(\d+)\)?)?", text)
    if not match:
        return ""
    target_text = match.group(0)
    return expr_for_existing_options(current_qid, expand_codes(target_text), all_vars)


def resolve_external_qid(token: str, all_vars: set[str]) -> str:
    """Resolve a questionnaire qid without silently accepting an ambiguous typo."""
    normalized = norm_qid(token)
    candidates = {normalized}
    # In Word questionnaires, a zero is occasionally typed as a capital O. Only
    # use this correction when it produces exactly one real question in ``all``.
    if "O" in normalized:
        candidates.add(normalized.replace("O", "0"))
    matches = [qid for qid in candidates if target_vars_for_qid(qid, all_vars)]
    return matches[0] if len(matches) == 1 else ""


def question_option_codes(docx_path: Path, all_vars: set[str]) -> dict[str, list[str]]:
    """Collect option codes from question blocks for scalar "all options" mutexes."""
    doc = Document(docx_path)
    codes: dict[str, set[str]] = {}
    current_qid = ""
    for paragraph in doc.paragraphs:
        text = unicodedata.normalize("NFKC", paragraph.text)
        qid = qid_from_question(text)
        if qid:
            current_qid = qid
        option_code = option_code_from_text(text)
        if current_qid and option_code and target_vars_for_qid(current_qid, all_vars):
            codes.setdefault(current_qid, set()).add(option_code)
    return {qid: sorted(values, key=int) for qid, values in codes.items()}


def parse_external_option_target(
    text: str,
    all_vars: set[str],
    question_codes: dict[str, list[str]],
) -> str:
    match = re.search(r"([A-Z][A-Z0-9_]*)選項\(?(\d+)\)?", text)
    if match:
        qid = resolve_external_qid(match.group(1), all_vars)
        return expr_for_option(qid, match.group(2), all_vars) if qid else ""

    match = re.search(r"([A-Z][A-Z0-9_]*)", text)
    if not match:
        return ""
    qid = resolve_external_qid(match.group(1), all_vars)
    if not qid:
        return ""
    codes = question_codes.get(qid, [])
    return expr_for_existing_options(qid, codes, all_vars) if codes else ""


def clean_condition_text(text: str) -> str:
    text = re.sub(r"者$", "", text)
    text = text.strip(" ,，、;；")
    return text


def parse_mutex_bracket(
    bracket: str,
    current_qid: str,
    option_code: str,
    all_vars: set[str],
    question_codes: dict[str, list[str]],
) -> tuple[str, str, str]:
    raw = unicodedata.normalize("NFKC", bracket)
    text = normalize(raw)
    if "互斥無法點選" in text:
        return "", "", "display check, not option mutex"
    if "互斥" not in text:
        return "", "", "no mutex cue"

    current_option = expr_for_option(current_qid, option_code, all_vars)
    if not current_option:
        return "", "", "current option variable not found"

    if "選項互斥" in text and "與" not in text:
        condition_text = clean_condition_text(text.split("選項互斥", 1)[0])
        parsed = parse_condition(condition_text, all_vars)
        return (parsed[0], current_option, "") if parsed else ("", "", "cannot parse condition")

    if "與此選項互斥" in text:
        condition_text = clean_condition_text(text.split("與此選項互斥", 1)[0])
        parsed = parse_condition(condition_text, all_vars)
        return (parsed[0], current_option, "") if parsed else ("", "", "cannot parse condition")

    if "與選項" in text:
        condition_text = clean_condition_text(text.split("與選項", 1)[0])
        mutex = parse_option_range_target(text.split("與", 1)[1], current_qid, all_vars)
        if not mutex:
            return "", "", "mutex option target not found"
        if condition_text:
            parsed = parse_condition(condition_text, all_vars)
            return (parsed[0], mutex, "") if parsed else ("", "", "cannot parse condition")
        return current_option, mutex, ""

    if text.startswith("與"):
        target_text = text[1:].split("互斥", 1)[0]
        parsed = parse_condition(target_text, all_vars)
        if parsed:
            return parsed[0], current_option, ""
        target = parse_external_option_target(target_text, all_vars, question_codes)
        if target:
            return current_option, target, ""
        return "", "", "cannot parse mutex target"

    if "互斥" in text:
        condition_text = clean_condition_text(text.split("互斥", 1)[0].replace("與", ""))
        parsed = parse_condition(condition_text, all_vars)
        return (parsed[0], current_option, "") if parsed else ("", "", "cannot parse condition")

    return "", "", "unsupported mutex wording"


def collect_mutex_rules(docx_path: Path, all_vars: set[str], qid_filter: str = "") -> tuple[list[MutexRule], list[dict[str, str]]]:
    current_qid = ""
    rules: list[MutexRule] = []
    review: list[dict[str, str]] = []
    question_codes = question_option_codes(docx_path, all_vars)
    qid_filter = norm_qid(qid_filter) if qid_filter else ""
    for unit in text_units(docx_path, all_vars):
        normalized_source = unicodedata.normalize("NFKC", unit.text)
        option_code = option_code_from_text(normalized_source)
        if not option_code:
            # Table text must never change paragraph context. In particular, an
            # unresolved table must not inherit the preceding question's qid.
            maybe_qid = qid_from_question(normalized_source) if unit.table is None else ""
            if maybe_qid:
                current_qid = maybe_qid
            continue
        for bracket in re.findall(r"【([^】]+)】", normalized_source):
            if "互斥" not in bracket:
                continue
            table = unit.table
            if table is not None and not table.qid:
                review.append(
                    {
                        "qid": "",
                        "option_code": option_code,
                        "reason": table.reason,
                        "bracket": bracket,
                        "source": normalized_source,
                        "table_index": str(table.table_index),
                        "table_option_codes": ",".join(table.option_codes),
                        "candidate_qids": ",".join(table.candidates),
                    }
                )
                continue
            active_qid = table.qid if table is not None else current_qid
            if not active_qid:
                continue
            if qid_filter and active_qid != qid_filter:
                continue
            condition, mutex, reason = parse_mutex_bracket(
                bracket,
                active_qid,
                option_code,
                all_vars,
                question_codes,
            )
            if condition and mutex:
                if same_question_mutex(active_qid, condition, mutex):
                    continue
                rules.append(
                    MutexRule(
                        qid=active_qid,
                        option_code=option_code,
                        condition=condition,
                        mutex=mutex,
                        bracket=bracket,
                        source_text=normalized_source,
                    )
                )
            elif reason != "display check, not option mutex":
                review.append(
                    {
                        "qid": active_qid,
                        "option_code": option_code,
                        "reason": reason,
                        "bracket": bracket,
                        "source": normalized_source,
                    }
                )
    return rules, review


def expr_vars(expr: str) -> set[str]:
    return set(re.findall(r"\bv[A-Za-z_][A-Za-z0-9_]*\b", expr))


def var_belongs_to_qid(var_name: str, qid: str) -> bool:
    base = re.escape(f"v{norm_qid(qid)}")
    # ``vZA0_1`` is a separate question from ``vZA0``. Only recognised grouped
    # suffixes belong to the same question, so the same-question guard cannot
    # suppress a genuine cross-question mutex rule.
    return bool(
        re.fullmatch(
            rf"{base}(?:m\d+|g\d+|s[A-Za-z0-9_]*|city|town)?",
            var_name,
            flags=re.IGNORECASE,
        )
    )


def same_question_mutex(qid: str, condition: str, mutex: str) -> bool:
    vars_found = expr_vars(condition) | expr_vars(mutex)
    return bool(vars_found) and all(var_belongs_to_qid(var_name, qid) for var_name in vars_found)


def ensure_mutex_header(ws: openpyxl.worksheet.worksheet.Worksheet) -> dict[str, int]:
    headers = {cell_text(ws.cell(1, col).value): col for col in range(1, ws.max_column + 1)}
    if MUTEX_HEADER not in headers:
        if "限制" not in headers:
            raise ValueError("邏輯組 missing 限制 column")
        insert_at = headers["限制"] + 1
        ws.insert_cols(insert_at)
        ws.cell(1, insert_at).value = MUTEX_HEADER
        headers = {cell_text(ws.cell(1, col).value): col for col in range(1, ws.max_column + 1)}
    for required in ("m", "p", "條件", "應答", "不應答", "限制", MUTEX_HEADER):
        if required not in headers:
            raise ValueError(f"邏輯組 missing {required} column")
    return headers


def existing_keys(ws: openpyxl.worksheet.worksheet.Worksheet, headers: dict[str, int]) -> set[tuple[str, str]]:
    keys: set[tuple[str, str]] = set()
    for row in range(2, ws.max_row + 1):
        keys.add((cell_text(ws.cell(row, headers["條件"]).value), cell_text(ws.cell(row, headers[MUTEX_HEADER]).value)))
    return keys


def write_csv(path: Path, rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = [
        "status", "row", "qid", "option_code", "條件", MUTEX_HEADER, "reason",
        "table_index", "table_option_codes", "candidate_qids", "bracket", "source",
    ]
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def refresh_mutex(docx_path: Path, workbook_path: Path, *, apply: bool, qid_filter: str = "", csv_path: Path | None = None, report_path: Path | None = None) -> dict[str, object]:
    all_vars = load_all_vars(workbook_path)
    rules, review = collect_mutex_rules(docx_path, all_vars, qid_filter)
    wb = openpyxl.load_workbook(workbook_path)
    ws = wb[LOGIC_SHEET]
    headers = ensure_mutex_header(ws)
    keys = existing_keys(ws, headers)
    output_rows: list[dict[str, str]] = []
    appended = 0
    duplicates = 0
    for rule in rules:
        key = (rule.condition, rule.mutex)
        row_data = {
            "qid": rule.qid,
            "option_code": rule.option_code,
            "條件": rule.condition,
            MUTEX_HEADER: rule.mutex,
            "bracket": rule.bracket,
            "source": rule.source_text,
        }
        if key in keys:
            duplicates += 1
            output_rows.append({"status": "duplicate", "row": "", **row_data})
            continue
        if apply:
            row_idx = ws.max_row + 1
            ws.cell(row_idx, headers["條件"]).value = rule.condition
            ws.cell(row_idx, headers[MUTEX_HEADER]).value = rule.mutex
            output_rows.append({"status": "appended", "row": str(row_idx), **row_data})
            keys.add(key)
            appended += 1
        else:
            output_rows.append({"status": "proposal", "row": "", **row_data})
    for item in review:
        output_rows.append({"status": "review", "row": "", "條件": "", MUTEX_HEADER: "", **item})

    if apply and (appended or MUTEX_HEADER in headers):
        backup = workbook_path.parent / "generated" / f"{workbook_path.stem}.before_logic_mutex_refresh.xlsx"
        backup.parent.mkdir(parents=True, exist_ok=True)
        if not backup.exists():
            shutil.copy2(workbook_path, backup)
        wb.save(workbook_path)

    if csv_path:
        write_csv(csv_path, output_rows)
    report = {
        "docx": str(docx_path),
        "workbook": str(workbook_path),
        "apply": apply,
        "qid_filter": norm_qid(qid_filter) if qid_filter else "",
        "rules": len(rules),
        "appended": appended,
        "duplicates": duplicates,
        "review": len(review),
    }
    if report_path:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Refresh option mutex rules from questionnaire Word into 邏輯組.互斥.")
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--workbook", required=True, type=Path)
    parser.add_argument("--apply", action="store_true")
    parser.add_argument("--qid", default="")
    parser.add_argument("--proposal-csv", type=Path)
    parser.add_argument("--report", type=Path)
    args = parser.parse_args()
    try:
        report = refresh_mutex(args.docx, args.workbook, apply=args.apply, qid_filter=args.qid, csv_path=args.proposal_csv, report_path=args.report)
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise
    print(json.dumps(report, ensure_ascii=False))


if __name__ == "__main__":
    main()
