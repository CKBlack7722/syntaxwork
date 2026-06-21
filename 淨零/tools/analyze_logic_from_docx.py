from __future__ import annotations

import argparse
import csv
import json
import re
import unicodedata
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

import openpyxl
from docx import Document


QUESTION_RE = re.compile(r"^([A-Za-z]{1,5}[0-9][A-Za-z0-9_]*|CK[0-9A-Za-z_]+|Y[0-9][A-Za-z0-9_]*|T[0-9][A-Za-z0-9_]*)")
BRACKET_RE = re.compile(r"【([^】]*(?:跳至|跳到|跳答|顯示|不顯示|續問|不問|答|若)[^】]*)】")
OPTION_RE = re.compile(r"^（?0*([0-9]{1,3}|97|98|96)）?(.+)")


@dataclass
class LogicCue:
    source: str
    index: int
    qid: str
    cue_type: str
    option_value: str
    condition_text: str
    target_text: str
    raw_text: str
    status: str
    note: str


def cell_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def normalize_text(text: str) -> str:
    return unicodedata.normalize("NFKC", text or "").strip()


def iter_doc_lines(docx_path: Path) -> Iterable[tuple[str, int, str]]:
    doc = Document(docx_path)
    for index, para in enumerate(doc.paragraphs, start=1):
        text = normalize_text(para.text)
        if text:
            yield "paragraph", index, text
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                text = normalize_text(cell.text.strip().replace("\n", " / "))
                if text:
                    yield f"table{table_index}.r{row_index}.c{cell_index}", row_index, text


def parse_qid(text: str, previous_qid: str) -> str:
    normalized = normalize_text(text)
    match = QUESTION_RE.match(normalized)
    if not match:
        match = re.search(r"(?:^|】| |\t)([A-Za-z]{1,5}[0-9][A-Za-z0-9_]*|CK[0-9A-Za-z_]+|Y[0-9][A-Za-z0-9_]*|T[0-9][A-Za-z0-9_]*)[.‧．、 ]", normalized)
    if match:
        return match.group(1)
    return previous_qid


def parse_option_value(text: str) -> str:
    match = OPTION_RE.match(normalize_text(text))
    return match.group(1) if match else ""


def classify_cue(cue: str) -> str:
    if any(text in cue for text in ("不需回答", "不需答", "不須回答", "免答")):
        return "skip"
    if any(text in cue for text in ("需回答", "需答", "才需回答", "才需答")):
        return "need"
    if "互斥" in cue:
        return "mutex"
    if "跳" in cue:
        return "skip"
    if "顯示" in cue:
        return "show"
    if "不顯示" in cue:
        return "hide"
    if "續問" in cue or "不問" in cue:
        return "flow"
    return "condition"


def extract_target(cue: str) -> str:
    match = re.search(r"(?:跳至|跳到|顯示|不顯示|續問|不問)(.+)$", cue)
    return match.group(1).strip() if match else ""


def extract_cues(docx_path: Path) -> list[LogicCue]:
    cues: list[LogicCue] = []
    current_qid = ""
    table_qids: dict[str, str] = {}
    for source, index, text in iter_doc_lines(docx_path):
        is_table = source.startswith("table")
        if is_table:
            table_id = source.split(".", 1)[0]
            active_qid = parse_qid(text, table_qids.get(table_id, ""))
            table_qids[table_id] = active_qid
        else:
            current_qid = parse_qid(text, current_qid)
            active_qid = current_qid
        for match in BRACKET_RE.finditer(text):
            cue = match.group(1).strip()
            cue_type = classify_cue(cue)
            option_value = parse_option_value(text)
            status = "review"
            note = ""
            if cue_type in {"show", "skip", "need", "mutex"} and active_qid:
                status = "candidate"
            elif is_table and cue_type in {"show", "skip", "need", "mutex", "flow", "condition"}:
                note = "Table question context is unresolved; do not infer it from the preceding paragraph."
            if "顯示題" in cue or "不顯示題號" in cue:
                status = "ignore"
                note = "Display-only marker, not a respondent-answer logic rule."
            cues.append(
                LogicCue(
                    source=source,
                    index=index,
                    qid=active_qid,
                    cue_type=cue_type,
                    option_value=option_value,
                    condition_text=cue,
                    target_text=extract_target(cue),
                    raw_text=text,
                    status=status,
                    note=note,
                )
            )
    return cues


def workbook_logic_summary(workbook_path: Path) -> dict[str, object]:
    wb = openpyxl.load_workbook(workbook_path, read_only=True, data_only=True)
    logic_name = "邏輯組" if "邏輯組" in wb.sheetnames else ""
    if not logic_name:
        return {"has_logic_sheet": False}
    ws = wb[logic_name]
    rows = list(ws.iter_rows(values_only=True))
    headers = [cell_text(value) for value in rows[0]]
    filled = 0
    blank = 0
    for row in rows[1:]:
        payload = [cell_text(value) for value in row[2:6]]
        if any(payload):
            filled += 1
        else:
            blank += 1
    return {
        "has_logic_sheet": True,
        "sheet": logic_name,
        "headers": headers,
        "rows": max(ws.max_row - 1, 0),
        "filled_logic_rows": filled,
        "blank_logic_rows": blank,
    }


def write_csv(path: Path, rows: list[LogicCue]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(asdict(rows[0]).keys()) if rows else ["source"])
        writer.writeheader()
        for row in rows:
            writer.writerow(asdict(row))


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract candidate logic cues from a questionnaire Word file.")
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--workbook", required=True, type=Path)
    parser.add_argument("--output-csv", required=True, type=Path)
    parser.add_argument("--report", required=True, type=Path)
    args = parser.parse_args()

    cues = extract_cues(args.docx)
    write_csv(args.output_csv, cues)
    counts: dict[str, int] = {}
    for cue in cues:
        key = f"{cue.status}:{cue.cue_type}"
        counts[key] = counts.get(key, 0) + 1
    report = {
        "docx": str(args.docx),
        "workbook": str(args.workbook),
        "cue_count": len(cues),
        "counts": counts,
        "workbook_logic": workbook_logic_summary(args.workbook),
        "review_samples": [asdict(cue) for cue in cues if cue.status != "ignore"][:50],
    }
    args.report.parent.mkdir(parents=True, exist_ok=True)
    args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"cue_count": len(cues), "counts": counts, "workbook_logic": report["workbook_logic"]}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
