from __future__ import annotations

import argparse
import csv
import json
import re
import shutil
import unicodedata
from dataclasses import dataclass
from pathlib import Path

import openpyxl
from docx import Document
from openpyxl.styles import Alignment, Font, PatternFill

from analyze_logic_from_docx import extract_cues


LOGIC_SHEET = "邏輯組"
ALL_SHEET = "all"
REVIEW_SHEET = "邏輯組_人工確認"
CUE_SHEET = "邏輯組_問卷線索"
NEED_CUES = ("需回答此題", "需答此題", "才需回答此題", "才需答此題")
SKIP_CUES = ("不需回答此題", "不需答此題", "不須回答此題", "免答此題")
VERB_PATTERN = r"(有?回答|皆回答|有?答|皆答|未答|回答|選擇|選|為)"


@dataclass(frozen=True)
class Atom:
    expr: str
    inverse: str


@dataclass(frozen=True)
class Node:
    op: str
    children: tuple["Node | Atom", ...]


@dataclass(frozen=True)
class Rule:
    qid: str
    cue: str
    condition: str
    inverse_condition: str
    target_vars: tuple[str, ...]
    source_text: str


REVIEW_HEADERS = [
    "規則鍵",
    "處理狀態",
    "題號",
    "問卷原始文字",
    "指示類型",
    "條件成立（轉譯）",
    "條件成立時應答",
    "條件成立時不應答",
    "條件不成立（反向轉譯）",
    "條件不成立時應答",
    "條件不成立時不應答",
    "邏輯組直接列",
    "邏輯組反向列",
    "人工確認",
    "人工備註",
    "判讀/未解析原因",
]


def cell_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def normalize_text(value: str) -> str:
    text = unicodedata.normalize("NFKC", value or "")
    text = text.replace("～", "~").replace("−", "-")
    text = re.sub(r"\s+", "", text)
    return text


def norm_qid(qid: str) -> str:
    qid = unicodedata.normalize("NFKC", qid or "").upper().strip()

    def strip_zero(match: re.Match[str]) -> str:
        prefix, number = match.groups()
        return prefix + str(int(number))

    return re.sub(r"([A-Z]+)0+(\d+)", strip_zero, qid)


def doc_text_units(docx_path: Path) -> list[str]:
    doc = Document(docx_path)
    units: list[str] = []
    current_qid = ""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            clean = unicodedata.normalize("NFKC", text)
            match = re.match(r"\s*([A-Z][A-Z0-9_]{0,12})(?=[^A-Z0-9_]|$)", clean)
            if match:
                current_qid = norm_qid(match.group(1))
            units.append(text)
            normalized = normalize_text(text)
            has_logic_cue = any(cue in normalized for cue in (*SKIP_CUES, *NEED_CUES))
            if current_qid and has_logic_cue and not match:
                units.append(f"{current_qid} {text}")
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                units.append(" ".join(cells))
    return units


def load_all_vars(workbook_path: Path) -> set[str]:
    wb = openpyxl.load_workbook(workbook_path, data_only=True, read_only=True)
    ws = wb[ALL_SHEET]
    return {cell_text(row[0]) for row in ws.iter_rows(min_row=2, values_only=True) if cell_text(row[0])}


def logic_headers(ws: openpyxl.worksheet.worksheet.Worksheet) -> dict[str, int]:
    headers = {cell_text(ws.cell(1, col).value): col for col in range(1, ws.max_column + 1)}
    for required in ("m", "p", "條件", "應答", "不應答", "限制"):
        if required not in headers:
            raise ValueError(f"{LOGIC_SHEET} missing header: {required}")
    return headers


def existing_keys(ws: openpyxl.worksheet.worksheet.Worksheet, headers: dict[str, int]) -> set[tuple[str, str, str, str]]:
    keys: set[tuple[str, str, str, str]] = set()
    for row in range(2, ws.max_row + 1):
        keys.add(
            (
                cell_text(ws.cell(row, headers["條件"]).value),
                cell_text(ws.cell(row, headers["應答"]).value),
                cell_text(ws.cell(row, headers["不應答"]).value),
                cell_text(ws.cell(row, headers["限制"]).value),
            )
        )
    return keys


def qid_from_text(text: str) -> str:
    clean = unicodedata.normalize("NFKC", text.strip())
    clean = re.sub(r"^[\s*◎●○壹貳參一二三四五六七八九十、.()（）-]+", "", clean)
    match = re.match(r"([A-Z][A-Z0-9_]{0,12})(?=[^A-Z0-9_]|$)", clean)
    return norm_qid(match.group(1)) if match else ""


def find_all_var(name: str, all_vars: set[str]) -> str:
    if name in all_vars:
        return name
    lower = name.lower()
    for var_name in all_vars:
        if var_name.lower() == lower:
            return var_name
    return ""


def target_vars_for_qid(qid: str, all_vars: set[str]) -> tuple[str, ...]:
    qid = norm_qid(qid)
    base = f"v{qid}"
    exact = find_all_var(base, all_vars)
    if exact:
        return (exact,)
    pattern = re.compile(rf"^v{re.escape(qid)}(?:m\d+|s[A-Za-z0-9_]+|s\d+|g\d+|city|town)$", re.IGNORECASE)
    found = [name for name in all_vars if pattern.match(name)]
    return tuple(sorted(found, key=natural_key))


def natural_key(value: str) -> tuple[object, ...]:
    parts = re.split(r"(\d+)", value)
    return tuple(int(part) if part.isdigit() else part.lower() for part in parts)


def option_expr(qid: str, code: str, all_vars: set[str]) -> Node | Atom | None:
    qid = norm_qid(qid)
    code_num = str(int(code))
    code2 = f"{int(code):02d}"
    for var_name in (f"v{qid}m{code2}", f"v{qid}m{code_num}"):
        exact = find_all_var(var_name, all_vars)
        if exact:
            return Atom(f"{exact} in 1", f"{exact}~=1")
    base = f"v{qid}"
    exact = find_all_var(base, all_vars)
    if exact:
        return Atom(f"{exact} in {code_num}", f"{exact}~={code_num}")
    grouped = tuple(name for name in target_vars_for_qid(qid, all_vars) if re.match(rf"^v{re.escape(qid)}g\d+$", name, re.IGNORECASE))
    if grouped and code_num == "0":
        return combine("&", [Atom(f"{name} in 0", f"{name}~=0") for name in grouped])
    return None


def expand_codes(tail: str) -> list[str]:
    codes: list[str] = []
    for start, end in re.findall(r"\(?(\d+)\)?(?:\s*[~\-至]\s*\(?(\d+)\)?)?", tail):
        if end:
            a, b = int(start), int(end)
            step = 1 if a <= b else -1
            codes.extend(str(value) for value in range(a, b + step, step))
        else:
            codes.append(str(int(start)))
    return codes


def render_node(node: Node | Atom, parent_op: str = "") -> str:
    if isinstance(node, Atom):
        return node.expr
    joiner = f" {node.op} "
    text = joiner.join(render_node(child, node.op) for child in node.children)
    if parent_op and parent_op != node.op and len(node.children) > 1:
        return f"({text})"
    return text


def invert_node(node: Node | Atom, parent_op: str = "") -> str:
    if isinstance(node, Atom):
        return node.inverse
    inverted_op = "|" if node.op == "&" else "&"
    text = f" {inverted_op} ".join(invert_node(child, inverted_op) for child in node.children)
    if parent_op and parent_op != inverted_op and len(node.children) > 1:
        return f"({text})"
    return text


def combine(op: str, children: list[Node | Atom]) -> Node | Atom | None:
    compact = [child for child in children if child]
    if not compact:
        return None
    if len(compact) == 1:
        return compact[0]
    return Node(op, tuple(compact))


def parse_condition_atom(part: str, all_vars: set[str]) -> Node | Atom | None:
    part = normalize_text(part)
    part = re.sub(r"者$", "", part)
    part = re.sub(r"([A-Z][A-Z0-9_]*)" + VERB_PATTERN + r"\((\d+)\)[、,]?或\((\d+)\)", r"\1\2(\3)或\1\2(\4)", part)
    or_nodes: list[Node | Atom] = []
    qid_or_before_verb = re.search(r"^[A-Z][A-Z0-9_]*(?:或[A-Z][A-Z0-9_]*)+" + VERB_PATTERN, part)
    or_parts = [part] if qid_or_before_verb else re.split(r"或", part)
    for or_part in or_parts:
        or_part = or_part.strip(" ,，、")
        if not or_part:
            continue
        match = re.search(r"([A-Z][A-Z0-9_]*(?:[、,，及和或][A-Z][A-Z0-9_]*)*)" + VERB_PATTERN + r"(.+)$", or_part)
        if not match:
            return None
        qid_text = match.group(1)
        qids = [norm_qid(q) for q in re.findall(r"[A-Z][A-Z0-9_]*", qid_text)]
        verb = match.group(2)
        codes = expand_codes(match.group(3))
        if not qids or not codes:
            return None
        qid_nodes: list[Node | Atom] = []
        for qid in qids:
            code_atoms = [option_expr(qid, code, all_vars) for code in codes]
            found_atoms = [atom for atom in code_atoms if atom]
            if not found_atoms:
                return None
            if len(found_atoms) != len(code_atoms) and len(codes) <= 10:
                return None
            qid_node = combine("|", found_atoms)
            if qid_node and verb == "未答":
                qid_node = Atom(invert_node(qid_node), render_node(qid_node))
            qid_nodes.append(qid_node)
        qids_relation = "|" if "或" in qid_text else "&"
        qid_node = combine(qids_relation, [node for node in qid_nodes if node])
        if qid_node:
            or_nodes.append(qid_node)
    return combine("|", or_nodes)


def parse_condition(condition_text: str, all_vars: set[str]) -> tuple[str, str] | None:
    text = normalize_text(condition_text)
    text = text.strip("【】[]()（）,，;；。")
    if not text:
        return None
    and_nodes: list[Node | Atom] = []
    for part in re.split(r"且|並且|以及", text):
        part = part.strip(" ,，、")
        if not part:
            continue
        node = parse_condition_atom(part, all_vars)
        if node is None:
            return None
        and_nodes.append(node)
    root = combine("&", and_nodes)
    if root is None:
        return None
    return render_node(root), invert_node(root)


def grouped_expr(exprs: list[str], op: str) -> str:
    unique: list[str] = []
    for expr in exprs:
        if expr and expr not in unique:
            unique.append(expr)
    if len(unique) == 1:
        return unique[0]
    wrapped = [f"({expr})" if " | " in expr or " & " in expr else expr for expr in unique]
    return f" {op} ".join(wrapped)


def bracket_rules(text: str, all_vars: set[str]) -> tuple[list[Rule], list[dict[str, str]]]:
    source = unicodedata.normalize("NFKC", text)
    qid = qid_from_text(source)
    rules: list[Rule] = []
    review: list[dict[str, str]] = []
    if not qid:
        return rules, review
    targets = target_vars_for_qid(qid, all_vars)
    parsed_rules: list[tuple[str, str, str]] = []
    for bracket in re.findall(r"【([^】]+)】", source):
        normalized = normalize_text(bracket)
        cue = ""
        if any(skip in normalized for skip in SKIP_CUES):
            cue = "skip"
        elif any(need in normalized for need in NEED_CUES):
            cue = "need"
        else:
            continue
        if not targets:
            review.append({"qid": qid, "reason": "target variables not found in all sheet", "text": source, "bracket": bracket})
            continue
        condition_text = normalized
        for phrase in (*SKIP_CUES, *NEED_CUES):
            condition_text = condition_text.replace(phrase, "")
        condition_text = condition_text.replace("才", "")
        parsed = parse_condition(condition_text, all_vars)
        if not parsed:
            review.append({"qid": qid, "reason": "cannot parse condition", "text": source, "bracket": bracket})
            continue
        condition, inverse = parsed
        parsed_rules.append((cue, condition, inverse))
    for grouped_cue in ("need", "skip"):
        cue_rules = [item for item in parsed_rules if item[0] == grouped_cue]
        if not cue_rules:
            continue
        if grouped_cue == "need":
            condition = grouped_expr([item[1] for item in cue_rules], "&")
            inverse = grouped_expr([item[2] for item in cue_rules], "|")
        else:
            condition = grouped_expr([item[1] for item in cue_rules], "|")
            inverse = grouped_expr([item[2] for item in cue_rules], "&")
        rules.append(Rule(qid=qid, cue=grouped_cue, condition=condition, inverse_condition=inverse, target_vars=targets, source_text=source))
    return rules, review


def build_rows(rules: list[Rule]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for rule in rules:
        targets = ",".join(rule.target_vars)
        if rule.cue == "need":
            rows.append({"條件": rule.condition, "應答": targets, "不應答": "", "限制": "", "qid": rule.qid, "source": rule.source_text, "direction": "need"})
            rows.append({"條件": rule.inverse_condition, "應答": "", "不應答": targets, "限制": "", "qid": rule.qid, "source": rule.source_text, "direction": "need_inverse"})
        else:
            rows.append({"條件": rule.condition, "應答": "", "不應答": targets, "限制": "", "qid": rule.qid, "source": rule.source_text, "direction": "skip"})
            rows.append({"條件": rule.inverse_condition, "應答": targets, "不應答": "", "限制": "", "qid": rule.qid, "source": rule.source_text, "direction": "skip_inverse"})
    return rows


def write_csv(path: Path, rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = ["status", "row", "qid", "direction", "條件", "應答", "不應答", "限制", "source", "reason", "bracket", "text"]
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def review_key(qid: str, cue: str, condition: str, targets: tuple[str, ...]) -> str:
    return "|".join((qid, cue, condition, ",".join(targets)))


def save_previous_review_values(wb: openpyxl.Workbook) -> dict[str, tuple[str, str]]:
    if REVIEW_SHEET not in wb.sheetnames:
        return {}
    ws = wb[REVIEW_SHEET]
    headers = {cell_text(cell.value): cell.column for cell in ws[1] if cell_text(cell.value)}
    key_col = headers.get("規則鍵")
    confirmation_col = headers.get("人工確認")
    note_col = headers.get("人工備註")
    saved: dict[str, tuple[str, str]] = {}
    if key_col:
        for row_idx in range(2, ws.max_row + 1):
            key = cell_text(ws.cell(row_idx, key_col).value)
            if key:
                saved[key] = (
                    cell_text(ws.cell(row_idx, confirmation_col).value) if confirmation_col else "",
                    cell_text(ws.cell(row_idx, note_col).value) if note_col else "",
                )
    wb.remove(ws)
    return saved


def logic_row_lookup(rows: list[dict[str, str]]) -> dict[tuple[str, str, str], tuple[str, str]]:
    lookup: dict[tuple[str, str, str], tuple[str, str]] = {}
    for row in rows:
        direction = cell_text(row.get("direction"))
        if direction not in {"need", "need_inverse", "skip", "skip_inverse"}:
            continue
        key = (cell_text(row.get("qid")), direction, cell_text(row.get("source")))
        lookup[key] = (cell_text(row.get("status")), cell_text(row.get("row")))
    return lookup


def logic_cells_for_rule(rule: Rule) -> tuple[str, str, str, str, str, str]:
    targets = ",".join(rule.target_vars)
    if rule.cue == "need":
        return targets, "", "", targets, "need", "need_inverse"
    return "", targets, targets, "", "skip", "skip_inverse"


def write_review_sheet(wb: openpyxl.Workbook, rules: list[Rule], review: list[dict[str, str]], output_rows: list[dict[str, str]]) -> None:
    """Write one readable row per questionnaire rule, preserving user decisions."""
    saved_values = save_previous_review_values(wb)
    ws = wb.create_sheet(REVIEW_SHEET)
    ws.append(REVIEW_HEADERS)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:P1"
    ws.row_dimensions[1].height = 30

    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    lookup = logic_row_lookup(output_rows)
    for rule in rules:
        direct_required, direct_forbidden, inverse_required, inverse_forbidden, direct_direction, inverse_direction = logic_cells_for_rule(rule)
        key = review_key(rule.qid, rule.cue, rule.condition, rule.target_vars)
        direct_status, direct_row = lookup.get((rule.qid, direct_direction, rule.source_text), ("", ""))
        inverse_status, inverse_row = lookup.get((rule.qid, inverse_direction, rule.source_text), ("", ""))
        status = direct_status or inverse_status or "proposal"
        confirmation, note = saved_values.get(key, ("", ""))
        ws.append(
            [
                key,
                status,
                rule.qid,
                rule.source_text,
                "需回答" if rule.cue == "need" else "不需回答",
                rule.condition,
                direct_required,
                direct_forbidden,
                rule.inverse_condition,
                inverse_required,
                inverse_forbidden,
                direct_row,
                inverse_row,
                confirmation,
                note,
                "",
            ]
        )

    for item in review:
        qid = cell_text(item.get("qid"))
        reason = cell_text(item.get("reason"))
        source = cell_text(item.get("text")) or cell_text(item.get("source"))
        bracket = cell_text(item.get("bracket"))
        key = review_key(qid, "review", bracket or source, ())
        confirmation, note = saved_values.get(key, ("", ""))
        ws.append(
            [
                key,
                "需人工判讀",
                qid,
                source,
                "未解析",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                "",
                confirmation,
                note,
                reason + (f" | 標記：{bracket}" if bracket else ""),
            ]
        )

    widths = {
        "A": 44, "B": 14, "C": 12, "D": 62, "E": 12, "F": 36, "G": 28, "H": 28,
        "I": 36, "J": 28, "K": 28, "L": 14, "M": 14, "N": 14, "O": 28, "P": 42,
    }
    status_fills = {
        "appended": PatternFill("solid", fgColor="E2F0D9"),
        "duplicate": PatternFill("solid", fgColor="DDEBF7"),
        "proposal": PatternFill("solid", fgColor="FFF2CC"),
        "需人工判讀": PatternFill("solid", fgColor="FCE4D6"),
    }
    for column, width in widths.items():
        ws.column_dimensions[column].width = width
    for row_idx in range(2, ws.max_row + 1):
        status = cell_text(ws.cell(row_idx, 2).value)
        fill = status_fills.get(status)
        for cell in ws[row_idx]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if fill:
                cell.fill = fill


def write_cue_sheet(wb: openpyxl.Workbook, docx_path: Path, rules: list[Rule], review: list[dict[str, str]]) -> None:
    """List every detected questionnaire cue so missing coverage is reviewable."""
    saved_values: dict[str, tuple[str, str]] = {}
    if CUE_SHEET in wb.sheetnames:
        previous = wb[CUE_SHEET]
        previous_headers = {cell_text(cell.value): cell.column for cell in previous[1] if cell_text(cell.value)}
        key_col = previous_headers.get("線索鍵")
        confirmation_col = previous_headers.get("人工確認")
        note_col = previous_headers.get("人工備註")
        if key_col:
            for row_idx in range(2, previous.max_row + 1):
                key = cell_text(previous.cell(row_idx, key_col).value)
                if key:
                    saved_values[key] = (
                        cell_text(previous.cell(row_idx, confirmation_col).value) if confirmation_col else "",
                        cell_text(previous.cell(row_idx, note_col).value) if note_col else "",
                    )
        wb.remove(previous)

    ws = wb.create_sheet(CUE_SHEET)
    headers = ["線索鍵", "題號", "來源", "位置", "線索類型", "問卷邏輯文字", "完整原始文字", "自動對照結果", "對照規則數", "人工確認", "人工備註"]
    ws.append(headers)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = "A1:K1"
    header_fill = PatternFill("solid", fgColor="548235")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    matched_source_counts: dict[str, int] = {}
    for rule in rules:
        matched_source_counts[rule.source_text] = matched_source_counts.get(rule.source_text, 0) + 1
    review_sources = {cell_text(item.get("text")) or cell_text(item.get("source")) for item in review}
    status_fills = {
        "已產生對照規則": PatternFill("solid", fgColor="E2F0D9"),
        "需人工判讀": PatternFill("solid", fgColor="FCE4D6"),
        "未對照：請檢查": PatternFill("solid", fgColor="FFF2CC"),
        "忽略：非資料應答邏輯": PatternFill("solid", fgColor="D9EAF7"),
    }
    for cue in extract_cues(docx_path):
        key = f"{cue.source}|{cue.index}|{cue.condition_text}"
        matches = matched_source_counts.get(cue.raw_text, 0)
        if cue.status == "ignore":
            result = "忽略：非資料應答邏輯"
        elif matches:
            result = "已產生對照規則"
        elif cue.raw_text in review_sources:
            result = "需人工判讀"
        else:
            result = "未對照：請檢查"
        confirmation, note = saved_values.get(key, ("", ""))
        ws.append([
            key, cue.qid, cue.source, cue.index, cue.cue_type, cue.condition_text,
            cue.raw_text, result, matches, confirmation, note,
        ])
        fill = status_fills.get(result)
        if fill:
            for cell in ws[ws.max_row]:
                cell.fill = fill

    for column, width in {"A": 36, "B": 12, "C": 14, "D": 9, "E": 14, "F": 44, "G": 70, "H": 26, "I": 13, "J": 14, "K": 28}.items():
        ws.column_dimensions[column].width = width
    for row_idx in range(2, ws.max_row + 1):
        for cell in ws[row_idx]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)



def refresh_logic(docx_path: Path, workbook_path: Path, *, apply: bool, qid_filter: str = "", csv_path: Path | None = None, report_path: Path | None = None) -> dict[str, object]:
    all_vars = load_all_vars(workbook_path)
    units = doc_text_units(docx_path)
    rules: list[Rule] = []
    review: list[dict[str, str]] = []
    for unit in units:
        unit_rules, unit_review = bracket_rules(unit, all_vars)
        rules.extend(unit_rules)
        review.extend(unit_review)
    if qid_filter:
        qid_filter = norm_qid(qid_filter)
        rules = [rule for rule in rules if rule.qid == qid_filter]
        review = [item for item in review if item.get("qid") == qid_filter]

    proposed = build_rows(rules)
    wb = openpyxl.load_workbook(workbook_path)
    ws = wb[LOGIC_SHEET]
    headers = logic_headers(ws)
    keys = existing_keys(ws, headers)
    output_rows: list[dict[str, str]] = []
    appended = 0
    duplicates = 0
    for row in proposed:
        key = (row["條件"], row["應答"], row["不應答"], row["限制"])
        if key in keys:
            duplicates += 1
            output_rows.append({"status": "duplicate", "row": "", **row})
            continue
        if apply:
            next_row = ws.max_row + 1
            for col_name in ("條件", "應答", "不應答", "限制"):
                ws.cell(next_row, headers[col_name]).value = row[col_name]
            output_rows.append({"status": "appended", "row": str(next_row), **row})
            keys.add(key)
            appended += 1
        else:
            output_rows.append({"status": "proposal", "row": "", **row})
    for item in review:
        output_rows.append({"status": "review", "row": "", "direction": "", "條件": "", "應答": "", "不應答": "", "限制": "", **item})

    if apply:
        write_review_sheet(wb, rules, review, output_rows)
        write_cue_sheet(wb, docx_path, rules, review)
        backup = workbook_path.parent / "generated" / f"{workbook_path.stem}.before_logic_refresh.xlsx"
        backup.parent.mkdir(parents=True, exist_ok=True)
        if not backup.exists():
            shutil.copy2(workbook_path, backup)
        wb.save(workbook_path)

    if csv_path:
        write_csv(csv_path, output_rows)
    report = {
        "docx": str(docx_path),
        "workbook": str(workbook_path),
        "apply": apply,
        "qid_filter": qid_filter,
        "source_units": len(units),
        "rules": len(rules),
        "proposed_rows": len(proposed),
        "appended": appended,
        "duplicates": duplicates,
        "review": len(review),
        "review_sheet": REVIEW_SHEET if apply else "",
        "cue_sheet": CUE_SHEET if apply else "",
    }
    if report_path:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Refresh missing questionnaire-derived logic rows in the Excel 邏輯組 sheet.")
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--workbook", required=True, type=Path)
    parser.add_argument("--apply", action="store_true")
    parser.add_argument("--qid", default="")
    parser.add_argument("--proposal-csv", type=Path)
    parser.add_argument("--report", type=Path)
    parser.add_argument(
        "--preview-workbook",
        type=Path,
        help="Copy the source workbook here, fill only the copy, and add 邏輯組_人工確認 for review.",
    )
    args = parser.parse_args()
    target_workbook = args.workbook
    target_apply = args.apply
    if args.preview_workbook:
        args.preview_workbook.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.workbook, args.preview_workbook)
        target_workbook = args.preview_workbook
        target_apply = True
    report = refresh_logic(
        args.docx,
        target_workbook,
        apply=target_apply,
        qid_filter=args.qid,
        csv_path=args.proposal_csv,
        report_path=args.report,
    )
    if args.preview_workbook:
        report["source_workbook"] = str(args.workbook)
        report["preview_workbook"] = str(args.preview_workbook)
    print(json.dumps(report, ensure_ascii=False))


if __name__ == "__main__":
    main()
