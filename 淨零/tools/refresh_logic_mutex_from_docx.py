from __future__ import annotations

import argparse
import csv
import json
import re
import shutil
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path

import openpyxl
from docx import Document

from refresh_logic_group_from_docx import (
    cell_text,
    expand_codes,
    load_all_vars,
    natural_key,
    norm_qid,
    option_expr,
    parse_condition,
    render_node,
)


LOGIC_SHEET = "邏輯組"
MUTEX_HEADER = "互斥"


@dataclass(frozen=True)
class MutexRule:
    qid: str
    option_code: str
    condition: str
    mutex: str
    bracket: str
    source_text: str


def normalize(value: str) -> str:
    text = unicodedata.normalize("NFKC", value or "")
    text = text.replace("～", "~").replace("−", "-")
    text = re.sub(r"\s+", "", text)
    return text


def text_units(docx_path: Path) -> list[str]:
    doc = Document(docx_path)
    units: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            units.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                units.append(" ".join(cells))
    return units


def qid_from_question(text: str) -> str:
    clean = unicodedata.normalize("NFKC", text.strip())
    clean = re.sub(r"^[\s*◎●○□壹貳參一二三四五六七八九十、.()（）-]+", "", clean)
    match = re.match(r"(?:CK)?([A-Z][A-Z0-9_]{0,12})(?=[^A-Z0-9_]|$)", clean)
    return norm_qid(match.group(1)) if match else ""


def option_code_from_text(text: str) -> str:
    match = re.search(r"(?:^|[□\s])\((\d+)\)", unicodedata.normalize("NFKC", text))
    return str(int(match.group(1))) if match else ""


def expr_for_option(qid: str, code: str, all_vars: set[str]) -> str:
    node = option_expr(qid, code, all_vars)
    return render_node(node) if node else ""


def expr_for_existing_options(qid: str, codes: list[str], all_vars: set[str]) -> str:
    exprs: list[str] = []
    for code in codes:
        expr = expr_for_option(qid, code, all_vars)
        if expr and expr not in exprs:
            exprs.append(expr)
    return " | ".join(exprs)


def parse_option_range_target(text: str, current_qid: str, all_vars: set[str]) -> str:
    match = re.search(r"選項\(?(\d+)\)?(?:\s*[~\-至]\s*\(?(\d+)\)?)?", text)
    if not match:
        return ""
    target_text = match.group(0)
    return expr_for_existing_options(current_qid, expand_codes(target_text), all_vars)


def parse_external_option_target(text: str, all_vars: set[str]) -> str:
    match = re.search(r"([A-Z][A-Z0-9_]*)(?:選項)?\(?(\d+)\)?", text)
    if not match:
        return ""
    return expr_for_option(match.group(1), match.group(2), all_vars)


def clean_condition_text(text: str) -> str:
    text = re.sub(r"者$", "", text)
    text = text.strip(" ,，、;；")
    return text


def parse_mutex_bracket(bracket: str, current_qid: str, option_code: str, all_vars: set[str]) -> tuple[str, str, str]:
    raw = unicodedata.normalize("NFKC", bracket)
    text = normalize(raw)
    if "互斥無法點選" in text:
        return "", "", "display check, not option mutex"
    if "互斥" not in text:
        return "", "", "no mutex cue"

    current_option = expr_for_option(current_qid, option_code, all_vars)
    if not current_option:
        return "", "", "current option variable not found"

    if "選項互斥" in text and "與" not in text:
        condition_text = clean_condition_text(text.split("選項互斥", 1)[0])
        parsed = parse_condition(condition_text, all_vars)
        return (parsed[0], current_option, "") if parsed else ("", "", "cannot parse condition")

    if "與此選項互斥" in text:
        condition_text = clean_condition_text(text.split("與此選項互斥", 1)[0])
        parsed = parse_condition(condition_text, all_vars)
        return (parsed[0], current_option, "") if parsed else ("", "", "cannot parse condition")

    if "與選項" in text:
        condition_text = clean_condition_text(text.split("與選項", 1)[0])
        mutex = parse_option_range_target(text.split("與", 1)[1], current_qid, all_vars)
        if not mutex:
            return "", "", "mutex option target not found"
        if condition_text:
            parsed = parse_condition(condition_text, all_vars)
            return (parsed[0], mutex, "") if parsed else ("", "", "cannot parse condition")
        return current_option, mutex, ""

    if text.startswith("與"):
        target_text = text[1:].split("互斥", 1)[0]
        parsed = parse_condition(target_text, all_vars)
        if parsed:
            return parsed[0], current_option, ""
        target = parse_external_option_target(target_text, all_vars)
        if target:
            return current_option, target, ""
        return "", "", "cannot parse mutex target"

    if "互斥" in text:
        condition_text = clean_condition_text(text.split("互斥", 1)[0].replace("與", ""))
        parsed = parse_condition(condition_text, all_vars)
        return (parsed[0], current_option, "") if parsed else ("", "", "cannot parse condition")

    return "", "", "unsupported mutex wording"


def collect_mutex_rules(docx_path: Path, all_vars: set[str], qid_filter: str = "") -> tuple[list[MutexRule], list[dict[str, str]]]:
    current_qid = ""
    rules: list[MutexRule] = []
    review: list[dict[str, str]] = []
    qid_filter = norm_qid(qid_filter) if qid_filter else ""
    for source in text_units(docx_path):
        normalized_source = unicodedata.normalize("NFKC", source)
        option_code = option_code_from_text(normalized_source)
        if not option_code:
            maybe_qid = qid_from_question(normalized_source)
            if maybe_qid:
                current_qid = maybe_qid
            continue
        if not current_qid:
            continue
        if qid_filter and current_qid != qid_filter:
            continue
        for bracket in re.findall(r"【([^】]+)】", normalized_source):
            if "互斥" not in bracket:
                continue
            condition, mutex, reason = parse_mutex_bracket(bracket, current_qid, option_code, all_vars)
            if condition and mutex:
                if same_question_mutex(current_qid, condition, mutex):
                    continue
                rules.append(
                    MutexRule(
                        qid=current_qid,
                        option_code=option_code,
                        condition=condition,
                        mutex=mutex,
                        bracket=bracket,
                        source_text=normalized_source,
                    )
                )
            elif reason != "display check, not option mutex":
                review.append(
                    {
                        "qid": current_qid,
                        "option_code": option_code,
                        "reason": reason,
                        "bracket": bracket,
                        "source": normalized_source,
                    }
                )
    return rules, review


def expr_vars(expr: str) -> set[str]:
    return set(re.findall(r"\bv[A-Za-z_][A-Za-z0-9_]*\b", expr))


def var_belongs_to_qid(var_name: str, qid: str) -> bool:
    return var_name.lower().startswith(f"v{norm_qid(qid).lower()}")


def same_question_mutex(qid: str, condition: str, mutex: str) -> bool:
    vars_found = expr_vars(condition) | expr_vars(mutex)
    return bool(vars_found) and all(var_belongs_to_qid(var_name, qid) for var_name in vars_found)


def ensure_mutex_header(ws: openpyxl.worksheet.worksheet.Worksheet) -> dict[str, int]:
    headers = {cell_text(ws.cell(1, col).value): col for col in range(1, ws.max_column + 1)}
    if MUTEX_HEADER not in headers:
        if "限制" not in headers:
            raise ValueError("邏輯組 missing 限制 column")
        insert_at = headers["限制"] + 1
        ws.insert_cols(insert_at)
        ws.cell(1, insert_at).value = MUTEX_HEADER
        headers = {cell_text(ws.cell(1, col).value): col for col in range(1, ws.max_column + 1)}
    for required in ("m", "p", "條件", "應答", "不應答", "限制", MUTEX_HEADER):
        if required not in headers:
            raise ValueError(f"邏輯組 missing {required} column")
    return headers


def existing_keys(ws: openpyxl.worksheet.worksheet.Worksheet, headers: dict[str, int]) -> set[tuple[str, str]]:
    keys: set[tuple[str, str]] = set()
    for row in range(2, ws.max_row + 1):
        keys.add((cell_text(ws.cell(row, headers["條件"]).value), cell_text(ws.cell(row, headers[MUTEX_HEADER]).value)))
    return keys


def write_csv(path: Path, rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = ["status", "row", "qid", "option_code", "條件", MUTEX_HEADER, "reason", "bracket", "source"]
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def refresh_mutex(docx_path: Path, workbook_path: Path, *, apply: bool, qid_filter: str = "", csv_path: Path | None = None, report_path: Path | None = None) -> dict[str, object]:
    all_vars = load_all_vars(workbook_path)
    rules, review = collect_mutex_rules(docx_path, all_vars, qid_filter)
    wb = openpyxl.load_workbook(workbook_path)
    ws = wb[LOGIC_SHEET]
    headers = ensure_mutex_header(ws)
    keys = existing_keys(ws, headers)
    output_rows: list[dict[str, str]] = []
    appended = 0
    duplicates = 0
    for rule in rules:
        key = (rule.condition, rule.mutex)
        row_data = {
            "qid": rule.qid,
            "option_code": rule.option_code,
            "條件": rule.condition,
            MUTEX_HEADER: rule.mutex,
            "bracket": rule.bracket,
            "source": rule.source_text,
        }
        if key in keys:
            duplicates += 1
            output_rows.append({"status": "duplicate", "row": "", **row_data})
            continue
        if apply:
            row_idx = ws.max_row + 1
            ws.cell(row_idx, headers["條件"]).value = rule.condition
            ws.cell(row_idx, headers[MUTEX_HEADER]).value = rule.mutex
            output_rows.append({"status": "appended", "row": str(row_idx), **row_data})
            keys.add(key)
            appended += 1
        else:
            output_rows.append({"status": "proposal", "row": "", **row_data})
    for item in review:
        output_rows.append({"status": "review", "row": "", "條件": "", MUTEX_HEADER: "", **item})

    if apply and (appended or MUTEX_HEADER in headers):
        backup = workbook_path.parent / "generated" / f"{workbook_path.stem}.before_logic_mutex_refresh.xlsx"
        backup.parent.mkdir(parents=True, exist_ok=True)
        if not backup.exists():
            shutil.copy2(workbook_path, backup)
        wb.save(workbook_path)

    if csv_path:
        write_csv(csv_path, output_rows)
    report = {
        "docx": str(docx_path),
        "workbook": str(workbook_path),
        "apply": apply,
        "qid_filter": norm_qid(qid_filter) if qid_filter else "",
        "rules": len(rules),
        "appended": appended,
        "duplicates": duplicates,
        "review": len(review),
    }
    if report_path:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Refresh option mutex rules from questionnaire Word into 邏輯組.互斥.")
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--workbook", required=True, type=Path)
    parser.add_argument("--apply", action="store_true")
    parser.add_argument("--qid", default="")
    parser.add_argument("--proposal-csv", type=Path)
    parser.add_argument("--report", type=Path)
    args = parser.parse_args()
    try:
        report = refresh_mutex(args.docx, args.workbook, apply=args.apply, qid_filter=args.qid, csv_path=args.proposal_csv, report_path=args.report)
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise
    print(json.dumps(report, ensure_ascii=False))


if __name__ == "__main__":
    main()
